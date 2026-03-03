"""Services for exporting resumes and cover letters as DOCX/PDF."""

from enum import Enum
from io import BytesIO
from typing import Optional, Dict, Any
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY


class TemplateStyle(str, Enum):
    """Supported visual templates for generated documents."""

    ATS = "ats"
    MODERN = "modern"


class ExportFormat(str, Enum):
    """Supported export formats."""

    PDF = "pdf"
    DOCX = "docx"


class DocumentExportService:
    """Generate resume / cover letter documents from Markdown-like content."""

    def __init__(self) -> None:
        # Professional fonts and colors (defaults)
        self.base_font_name = "Helvetica"
        self.base_font_size = 11
        self.primary_color = HexColor("#2C3E50")  # Dark blue-gray
        self.accent_color = HexColor("#3498DB")   # Professional blue
        self.text_color = HexColor("#34495E")     # Readable gray

    def apply_template_styling(self, template_info: Optional[Dict[str, Any]]) -> None:
        """Apply template styling from extracted PDF template."""
        if not template_info:
            return
        
        # Apply fonts if available
        if template_info.get("fonts") and len(template_info["fonts"]) > 0:
            primary_font = template_info["fonts"][0]
            self.base_font_size = int(primary_font.get("size", 11))
        
        # Apply colors if available
        if template_info.get("colors") and len(template_info["colors"]) > 0:
            try:
                self.primary_color = HexColor(template_info["colors"][0])
                if len(template_info["colors"]) > 1:
                    self.accent_color = HexColor(template_info["colors"][1])
            except:
                pass  # Keep defaults if color parsing fails

    def generate_docx(
        self,
        *,
        title: str,
        content_markdown: str,
        full_name: Optional[str] = None,
        email: Optional[str] = None,
        template_style: TemplateStyle = TemplateStyle.ATS,
        document_type: Optional[str] = None,
    ) -> bytes:
        """Generate a professional DOCX document from Markdown-like content."""
        doc = Document()

        # Configure margins for professional look
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Configure base styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(self.base_font_size)

        header_name = full_name or title

        # Top header with name - Professional styling
        header_paragraph = doc.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_paragraph.add_run(header_name.upper())
        header_run.bold = True
        header_run.font.size = Pt(22 if template_style == TemplateStyle.MODERN else 20)
        header_run.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-gray

        # Optional contact line
        if email:
            contact_paragraph = doc.add_paragraph()
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_paragraph.add_run(email)
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = RGBColor(52, 73, 94)

        # Add horizontal line separator
        doc.add_paragraph("_" * 80)

        # Small spacer
        doc.add_paragraph()

        # Parse Markdown-like content into headings, bullets, and paragraphs
        lines = content_markdown.splitlines()
        in_list = False

        for raw_line in lines:
            line = raw_line.rstrip()
            stripped = line.strip()

            if not stripped:
                if in_list:
                    in_list = False
                continue

            # Skip duplicate name headers
            if stripped.startswith("# ") and stripped[2:].strip().upper() == header_name.upper():
                continue

            # Main headings (H1)
            if stripped.startswith("# "):
                heading_text = stripped[2:].strip()
                h1 = doc.add_heading(heading_text, level=1)
                h1.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                h1.runs[0].font.size = Pt(16)
                in_list = False
                continue

            # Section headings (H2)
            if stripped.startswith("## "):
                heading_text = stripped[3:].strip()
                h2 = doc.add_heading(heading_text, level=2)
                h2.runs[0].font.color.rgb = RGBColor(52, 152, 219)
                h2.runs[0].font.size = Pt(14)
                in_list = False
                continue

            # Subsection headings (H3)
            if stripped.startswith("### "):
                heading_text = stripped[4:].strip()
                h3 = doc.add_heading(heading_text, level=3)
                h3.runs[0].font.color.rgb = RGBColor(52, 73, 94)
                h3.runs[0].font.size = Pt(12)
                in_list = False
                continue

            # Bullet lists
            bullet_prefixes = ("- ", "* ", "• ", "✅ ")
            if any(stripped.startswith(prefix) for prefix in bullet_prefixes):
                for prefix in bullet_prefixes:
                    if stripped.startswith(prefix):
                        bullet_text = stripped[len(prefix):].strip()
                        break
                bullet_paragraph = doc.add_paragraph(style="List Bullet")
                bullet_paragraph.add_run(bullet_text)
                in_list = True
                continue

            # Fallback: regular paragraph
            if stripped:
                p = doc.add_paragraph(stripped)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                in_list = False

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_pdf(
        self,
        *,
        title: str,
        content_markdown: str,
        full_name: Optional[str] = None,
        email: Optional[str] = None,
        template_style: TemplateStyle = TemplateStyle.ATS,
        document_type: Optional[str] = None,
    ) -> bytes:
        """Generate a professional, well-formatted PDF from Markdown-like content."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=LETTER,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch,
        )

        # Container for the 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles for professional look
        name_style = ParagraphStyle(
            'CustomName',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=self.primary_color,
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
        )
        
        contact_style = ParagraphStyle(
            'CustomContact',
            parent=styles['Normal'],
            fontSize=10,
            textColor=self.text_color,
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica',
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=self.primary_color,
            spaceAfter=8,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            borderWidth=1,
            borderColor=self.accent_color,
            borderPadding=4,
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=self.accent_color,
            spaceAfter=6,
            spaceBefore=10,
            fontName='Helvetica-Bold',
        )
        
        heading3_style = ParagraphStyle(
            'CustomHeading3',
            parent=styles['Heading3'],
            fontSize=11,
            textColor=self.text_color,
            spaceAfter=4,
            spaceBefore=6,
            fontName='Helvetica-Bold',
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            textColor=self.text_color,
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            fontName='Helvetica',
            leading=14,
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontSize=10,
            textColor=self.text_color,
            spaceAfter=4,
            leftIndent=20,
            fontName='Helvetica',
            bulletIndent=10,
            leading=13,
        )

        # Header name
        header_name = full_name or title
        elements.append(Paragraph(header_name.upper(), name_style))
        
        # Contact info
        if email:
            elements.append(Paragraph(email, contact_style))
        
        # Horizontal line
        elements.append(Spacer(1, 0.1*inch))

        # Parse content
        lines = content_markdown.splitlines()
        
        for raw_line in lines:
            line = raw_line.rstrip()
            stripped = line.strip()

            if not stripped:
                elements.append(Spacer(1, 0.05*inch))
                continue

            # Skip duplicate name headers
            if stripped.startswith("# ") and stripped[2:].strip().upper() == header_name.upper():
                continue

            # Escape special characters for ReportLab
            stripped = stripped.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            # Main headings (H1)
            if stripped.startswith("# "):
                heading_text = stripped[2:].strip()
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph(heading_text.upper(), heading1_style))
                continue

            # Section headings (H2)
            if stripped.startswith("## "):
                heading_text = stripped[3:].strip()
                elements.append(Paragraph(heading_text, heading2_style))
                continue

            # Subsection headings (H3)
            if stripped.startswith("### "):
                heading_text = stripped[4:].strip()
                elements.append(Paragraph(heading_text, heading3_style))
                continue

            # Bullet lists
            bullet_prefixes = ("- ", "* ", "• ", "✅ ")
            is_bullet = False
            for prefix in bullet_prefixes:
                if stripped.startswith(prefix):
                    bullet_text = stripped[len(prefix):].strip()
                    elements.append(Paragraph(f"• {bullet_text}", bullet_style))
                    is_bullet = True
                    break
            
            if is_bullet:
                continue

            # Regular paragraph
            if stripped:
                elements.append(Paragraph(stripped, body_style))

        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()

